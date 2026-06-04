"""
通用文档解析器

支持多种文档格式的文本提取：
- PDF: 使用 pypdf 提取文本
- DOCX: 使用 python-docx 提取文本
- HTML: 使用 BeautifulSoup 提取纯文本
- TXT: 直接读取纯文本
- URL: 通过 HTTP 抓取并解析网页内容

环境依赖：
- pypdf (PDF 解析)
- python-docx (DOCX 解析)
- beautifulsoup4 (HTML 解析)
- httpx (URL 抓取)
"""

import os
import re
from typing import Dict, List, Optional
from dataclasses import dataclass, field

from src.unified_logger import default_logger as logger


@dataclass
class ParseResult:
    """文档解析结果"""
    text: str = ""
    title: str = ""
    file_type: str = ""
    page_count: int = 0
    metadata: Dict = field(default_factory=dict)
    success: bool = True
    error: Optional[str] = None


class DocumentParser:
    """
    通用文档解析器

    支持多种格式的文本提取，自动检测文件类型。
    所有解析方法均为异步，内部使用 run_in_executor 包装同步 I/O 操作。
    """

    # 支持的文件类型
    SUPPORTED_TYPES = {"pdf", "docx", "html", "htm", "txt", "md", "url"}

    # 文件扩展名到类型的映射
    EXTENSION_MAP = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".html": "html",
        ".htm": "html",
        ".txt": "txt",
        ".md": "txt",
        ".text": "txt",
        ".csv": "txt",
        ".json": "txt",
        ".xml": "html",
    }

    def __init__(self, max_file_size_mb: int = 50):
        """
        初始化文档解析器

        Args:
            max_file_size_mb: 最大文件大小限制（MB）
        """
        self._max_file_size = max_file_size_mb * 1024 * 1024

    async def parse(self, file_path: str, file_type: str = None) -> ParseResult:
        """
        解析文档为纯文本

        Args:
            file_path: 文件路径或 URL
            file_type: 文件类型（可选，自动检测）

        Returns:
            ParseResult 解析结果
        """
        # 自动检测文件类型
        if file_type is None:
            file_type = self._detect_type(file_path)

        file_type = file_type.lower().strip(".")

        if file_type not in self.SUPPORTED_TYPES:
            return ParseResult(
                success=False,
                error=f"不支持的文件类型: {file_type}，支持: {', '.join(self.SUPPORTED_TYPES)}",
                file_type=file_type,
            )

        try:
            if file_type == "url":
                return await self._parse_url(file_path)
            elif file_type == "pdf":
                return await self._parse_pdf(file_path)
            elif file_type == "docx":
                return await self._parse_docx(file_path)
            elif file_type in ("html", "htm"):
                return await self._parse_html_file(file_path)
            elif file_type in ("txt", "md", "text", "csv", "json", "xml"):
                return await self._parse_txt(file_path)
            else:
                return ParseResult(success=False, error=f"未实现的解析器: {file_type}")
        except Exception as e:
            logger.error(f"文档解析失败 [{file_path}]: {e}")
            return ParseResult(success=False, error=str(e), file_type=file_type)

    async def parse_text(self, text: str, content_type: str = "txt") -> ParseResult:
        """
        直接解析文本内容（无需文件）

        Args:
            text: 文本内容
            content_type: 内容类型 (txt/html)

        Returns:
            ParseResult 解析结果
        """
        try:
            if content_type == "html":
                clean_text = self._strip_html(text)
                return ParseResult(text=clean_text, file_type="html", page_count=1)
            else:
                return ParseResult(text=text, file_type="txt", page_count=1)
        except Exception as e:
            logger.error(f"文本解析失败: {e}")
            return ParseResult(success=False, error=str(e))

    def _detect_type(self, file_path: str) -> str:
        """自动检测文件类型"""
        # URL 检测
        if file_path.startswith(("http://", "https://")):
            return "url"

        # 扩展名检测
        _, ext = os.path.splitext(file_path.lower())
        return self.EXTENSION_MAP.get(ext, "txt")

    async def _parse_pdf(self, file_path: str) -> ParseResult:
        """解析 PDF 文件"""
        import asyncio

        def _read_pdf():
            try:
                from pypdf import PdfReader
            except ImportError:
                raise ImportError("请安装 pypdf: pip install pypdf")

            reader = PdfReader(file_path)
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())

            metadata = {}
            if reader.metadata:
                metadata = {
                    "author": reader.metadata.author or "",
                    "creator": reader.metadata.creator or "",
                    "producer": reader.metadata.producer or "",
                    "subject": reader.metadata.subject or "",
                }

            return pages, metadata

        try:
            loop = asyncio.get_event_loop()
            pages, metadata = await loop.run_in_executor(None, _read_pdf)

            full_text = "\n\n".join(pages)
            title = metadata.get("subject", "") or os.path.basename(file_path)

            return ParseResult(
                text=full_text,
                title=title,
                file_type="pdf",
                page_count=len(pages),
                metadata=metadata,
            )
        except ImportError:
            raise
        except Exception as e:
            logger.error(f"PDF 解析失败 [{file_path}]: {e}")
            return ParseResult(success=False, error=str(e), file_type="pdf")

    async def _parse_docx(self, file_path: str) -> ParseResult:
        """解析 DOCX 文件"""
        import asyncio

        def _read_docx():
            try:
                from docx import Document as DocxDocument
            except ImportError:
                raise ImportError("请安装 python-docx: pip install python-docx")

            doc = DocxDocument(file_path)

            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            metadata = {}
            if doc.core_properties:
                props = doc.core_properties
                metadata = {
                    "author": props.author or "",
                    "title": props.title or "",
                    "subject": props.subject or "",
                    "created": str(props.created) if props.created else "",
                    "modified": str(props.modified) if props.modified else "",
                }

            return paragraphs, metadata

        try:
            loop = asyncio.get_event_loop()
            paragraphs, metadata = await loop.run_in_executor(None, _read_docx)

            full_text = "\n\n".join(paragraphs)
            title = metadata.get("title", "") or os.path.basename(file_path)

            return ParseResult(
                text=full_text,
                title=title,
                file_type="docx",
                page_count=1,
                metadata=metadata,
            )
        except ImportError:
            raise
        except Exception as e:
            logger.error(f"DOCX 解析失败 [{file_path}]: {e}")
            return ParseResult(success=False, error=str(e), file_type="docx")

    async def _parse_html_file(self, file_path: str) -> ParseResult:
        """解析本地 HTML 文件"""
        import asyncio

        def _read_html():
            encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
            content = None
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                raise ValueError(f"无法解码文件: {file_path}")

            return content

        try:
            loop = asyncio.get_event_loop()
            html_content = await loop.run_in_executor(None, _read_html)
            return self._parse_html_content(html_content, source=file_path)
        except Exception as e:
            logger.error(f"HTML 文件解析失败 [{file_path}]: {e}")
            return ParseResult(success=False, error=str(e), file_type="html")

    async def _parse_url(self, url: str) -> ParseResult:
        """抓取并解析网页内容"""
        try:
            import httpx
        except ImportError:
            raise ImportError("请安装 httpx: pip install httpx")

        try:
            async with httpx.AsyncClient(
                timeout=30.0,
                follow_redirects=True,
                headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                    "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
                },
            ) as client:
                response = await client.get(url)
                response.raise_for_status()
                html_content = response.text

            result = self._parse_html_content(html_content, source=url)
            result.file_type = "url"
            result.metadata["url"] = url
            result.metadata["status_code"] = response.status_code
            return result
        except httpx.HTTPStatusError as e:
            logger.error(f"URL 抓取 HTTP 错误 [{url}]: {e.response.status_code}")
            return ParseResult(
                success=False,
                error=f"HTTP {e.response.status_code}: {e.response.reason_phrase}",
                file_type="url",
            )
        except Exception as e:
            logger.error(f"URL 抓取失败 [{url}]: {e}")
            return ParseResult(success=False, error=str(e), file_type="url")

    async def _parse_txt(self, file_path: str) -> ParseResult:
        """解析纯文本文件"""
        import asyncio

        def _read_txt():
            encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"无法解码文件: {file_path}")

        try:
            loop = asyncio.get_event_loop()
            content = await loop.run_in_executor(None, _read_txt)

            title = ""
            # 尝试从首行提取标题（Markdown 格式）
            first_line = content.split("\n", 1)[0].strip()
            if first_line.startswith("# "):
                title = first_line[2:].strip()

            return ParseResult(
                text=content,
                title=title or os.path.basename(file_path),
                file_type="txt",
                page_count=1,
            )
        except Exception as e:
            logger.error(f"TXT 解析失败 [{file_path}]: {e}")
            return ParseResult(success=False, error=str(e), file_type="txt")

    def _parse_html_content(self, html_content: str, source: str = "") -> ParseResult:
        """从 HTML 内容提取纯文本"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            # 降级方案：使用正则去除标签
            text = self._strip_html(html_content)
            return ParseResult(text=text, file_type="html", page_count=1)

        soup = BeautifulSoup(html_content, "html.parser")

        # 移除不需要的标签
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
            tag.decompose()

        # 提取标题
        title = ""
        title_tag = soup.find("title")
        if title_tag:
            title = title_tag.get_text(strip=True)
        elif soup.find("h1"):
            title = soup.find("h1").get_text(strip=True)

        # 提取正文（优先提取 article/main 标签内容）
        main_content = soup.find("article") or soup.find("main") or soup.find("body")
        if main_content:
            text = main_content.get_text(separator="\n", strip=True)
        else:
            text = soup.get_text(separator="\n", strip=True)

        # 清理多余空行
        text = re.sub(r"\n{3,}", "\n\n", text)

        return ParseResult(
            text=text,
            title=title or source,
            file_type="html",
            page_count=1,
            metadata={"source": source},
        )

    @staticmethod
    def _strip_html(html_content: str) -> str:
        """使用正则去除 HTML 标签（降级方案）"""
        # 去除 script 和 style
        text = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", html_content, flags=re.DOTALL | re.IGNORECASE)
        # 去除所有 HTML 标签
        text = re.sub(r"<[^>]+>", " ", text)
        # 解码常见 HTML 实体
        text = text.replace("&nbsp;", " ").replace("&", "&")
        text = text.replace("<", "<").replace(">", ">")
        text = text.replace("\u201c", '"').replace("\u201d", '"')
        text = text.replace("\u2018", "'").replace("\u2019", "'")
        # 清理空白
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def get_supported_types(self) -> List[str]:
        """获取支持的文件类型列表"""
        return sorted(self.SUPPORTED_TYPES)


# 全局单例
document_parser = DocumentParser()
